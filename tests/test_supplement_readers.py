from __future__ import annotations

import importlib
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas

from cbio_curation_assistant.cbioportal.classification import ClassificationResult
from cbio_curation_assistant.reports import curation as report_curation
from cbio_curation_assistant.supplements import readers
from cbio_curation_assistant.supplements.readers import (
    EmptySupplementaryFileError,
    MissingExternalReaderError,
    MissingReaderDependencyError,
    SupplementaryParseError,
    SupplementaryReadResult,
    SupplementaryReaderPreflightError,
    UnsupportedSupplementaryFormatError,
    discover_supplementary_files,
    read_supplementary_file,
    require_supplementary_reader_dependencies,
)


def classification_result(format_key: str = "CLINICAL_SAMPLE") -> ClassificationResult:
    return ClassificationResult(
        format_key=format_key,
        target_file="data_clinical_sample.txt",
        confidence=70.0,
        required_present=["patient_id", "sample_id"],
        required_missing=[],
        optional_present=[],
        detected_as_aliases={},
        all_scores=[],
        is_matrix=False,
        notes="",
        verdict="CLINICAL_SAMPLE",
        spec_source="embedded",
        spec_fetched_at="fixture",
    )


class SupplementReaderTest(unittest.TestCase):
    def test_csv_tsv_and_text_files_are_read_as_single_sheets(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            fixtures = {
                "table.csv": "sample,value\nS1,1\n",
                "table.tsv": "sample\tvalue\nS1\t1\n",
                "table.txt": "sample|value\nS1|1\n",
            }
            for filename, content in fixtures.items():
                with self.subTest(filename=filename):
                    path = root / filename
                    path.write_text(content, encoding="utf-8")

                    result = read_supplementary_file(path)

                    self.assertEqual(list(result.sheets), ["Sheet1"])
                    self.assertEqual(result.sheets["Sheet1"].shape, (2, 2))
                    self.assertEqual(result.warnings, ())

    def test_excel_reader_ignores_empty_sheets_and_blank_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "tables.xlsx"
            with pd.ExcelWriter(path) as writer:
                pd.DataFrame(
                    [[None, None], ["sample", "value"], ["S1", 1]]
                ).to_excel(
                    writer,
                    sheet_name="Data",
                    index=False,
                    header=False,
                )
                pd.DataFrame().to_excel(writer, sheet_name="Empty", index=False)

            result = read_supplementary_file(path)

            self.assertEqual(list(result.sheets), ["Data"])
            self.assertEqual(result.sheets["Data"].shape, (2, 2))

    def test_excel_reader_returns_readable_sheets_with_warnings(self) -> None:
        class PartiallyReadableWorkbook:
            sheet_names = ["Good", "Broken"]

            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return None

            def parse(self, sheet_name: str, *, header: int | None):
                self.assert_header(header)
                if sheet_name == "Broken":
                    raise ValueError("damaged worksheet")
                return pd.DataFrame([["sample", "value"], ["S1", "1"]])

            @staticmethod
            def assert_header(header: int | None) -> None:
                if header is not None:
                    raise AssertionError(f"Unexpected header: {header}")

        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "partial.xlsx"
            path.write_bytes(b"not-empty")
            with patch.object(
                readers.pd,
                "ExcelFile",
                return_value=PartiallyReadableWorkbook(),
            ):
                result = read_supplementary_file(path)

        self.assertEqual(list(result.sheets), ["Good"])
        self.assertEqual(len(result.warnings), 1)
        self.assertIn("damaged worksheet", result.warnings[0])

    def test_malformed_excel_file_has_explicit_parse_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "broken.xlsx"
            path.write_bytes(b"not an Excel workbook")

            with self.assertRaisesRegex(
                SupplementaryParseError,
                "Could not open Excel supplement",
            ):
                read_supplementary_file(path)

    def test_docx_reader_exposes_tables_and_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "supplement.docx"
            document = Document()
            document.add_paragraph("Study notes")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "sample"
            table.cell(0, 1).text = "value"
            table.cell(1, 0).text = "S1"
            table.cell(1, 1).text = "1"
            document.save(path)

            result = read_supplementary_file(path)

            self.assertEqual(set(result.sheets), {"Table_1", "Text"})
            self.assertEqual(result.sheets["Table_1"].shape, (2, 2))
            self.assertEqual(result.sheets["Text"].iloc[0, 0], "Study notes")

    def test_legacy_word_reader_requires_libreoffice(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "supplement.doc"
            path.write_bytes(b"legacy document")

            with (
                patch.object(readers.shutil, "which", return_value=None),
                self.assertRaisesRegex(
                    MissingExternalReaderError,
                    "libreoffice",
                ),
            ):
                read_supplementary_file(path)

    def test_pdf_reader_uses_pdfplumber_text_when_no_table_exists(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "supplement.pdf"
            pdf = canvas.Canvas(str(path))
            pdf.drawString(72, 720, "Sample S1 has a reported alteration")
            pdf.save()

            result = read_supplementary_file(path)

            self.assertIn("Text", result.sheets)
            rendered = " ".join(
                str(value) for value in result.sheets["Text"].iloc[:, 0]
            )
            self.assertIn("Sample S1", rendered)

    def test_pdf_reader_does_not_fall_back_when_pdfplumber_is_missing(self) -> None:
        real_import_module = importlib.import_module

        def import_without_pdfplumber(name: str):
            if name == "pdfplumber":
                raise ModuleNotFoundError("pdfplumber is unavailable")
            return real_import_module(name)

        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "supplement.pdf"
            path.write_bytes(b"%PDF-1.4\n")

            with (
                patch.object(
                    readers,
                    "import_module",
                    side_effect=import_without_pdfplumber,
                ),
                self.assertRaisesRegex(
                    MissingReaderDependencyError,
                    "require pdfplumber",
                ),
            ):
                read_supplementary_file(path)

    def test_unsupported_empty_and_archive_files_have_explicit_errors(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            unsupported = root / "table.data"
            unsupported.write_text("sample\tvalue\nS1\t1\n", encoding="utf-8")
            empty = root / "empty.csv"
            empty.touch()
            archive = root / "tables.zip"
            archive.write_bytes(b"not-empty")

            with self.assertRaises(UnsupportedSupplementaryFormatError):
                read_supplementary_file(unsupported)
            with self.assertRaises(EmptySupplementaryFileError):
                read_supplementary_file(empty)
            with self.assertRaisesRegex(
                UnsupportedSupplementaryFormatError,
                "must be extracted",
            ):
                read_supplementary_file(archive)

    def test_discovery_sorts_deduplicates_and_optionally_recurses(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            nested = root / "nested"
            nested.mkdir()
            first = root / "b.tsv"
            second = root / "a.csv"
            nested_file = nested / "c.xlsx"
            ignored = root / "notes.json"
            for path in (first, second, nested_file, ignored):
                path.write_text("fixture", encoding="utf-8")

            shallow = discover_supplementary_files([root, second])
            recursive = discover_supplementary_files([root], recursive=True)

        self.assertEqual(shallow, (second.resolve(), first.resolve()))
        self.assertEqual(
            recursive,
            (second.resolve(), first.resolve(), nested_file.resolve()),
        )

    def test_direct_unsupported_discovery_input_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "notes.json"
            path.write_text("{}", encoding="utf-8")

            with self.assertRaises(UnsupportedSupplementaryFormatError):
                discover_supplementary_files([path])

    def test_dependency_preflight_aggregates_required_readers(self) -> None:
        def missing_import(name: str):
            raise ModuleNotFoundError(f"{name} unavailable")

        with (
            patch.object(readers, "import_module", side_effect=missing_import),
            patch.object(readers.shutil, "which", return_value=None),
            self.assertRaises(SupplementaryReaderPreflightError) as raised,
        ):
            require_supplementary_reader_dependencies(
                ["supplement.doc", "supplement.pdf", "supplement.xlsx"]
            )

        dependencies = {issue.dependency for issue in raised.exception.issues}
        self.assertEqual(
            dependencies,
            {"libreoffice", "openpyxl", "pdfplumber", "python-docx"},
        )

    def test_supplement_analysis_returns_classification_records(self) -> None:
        dataframe = pd.DataFrame([["PATIENT_ID", "SAMPLE_ID"], ["P1", "S1"]])
        read_result = SupplementaryReadResult(
            path=Path("/tmp/source.xlsx"),
            sheets={"Clinical": dataframe},
            warnings=("Could not read sheet 'Notes': damaged worksheet",),
        )
        warnings: list[str] = []
        with (
            patch.object(
                report_curation,
                "read_supplementary_file",
                return_value=read_result,
            ),
            patch.object(
                report_curation.specification_sources,
                "get_embedded_spec",
                return_value={
                    "specs": [],
                    "source": "embedded",
                    "fetched_at": "fixture",
                    "version": "fixture",
                },
            ),
            patch.object(
                report_curation,
                "classify_sheet",
                return_value=classification_result(),
            ),
        ):
            records = report_curation.analyse_supplementary_files(
                ["/tmp/source.xlsx"],
                warnings=warnings,
            )

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].file, "source.xlsx")
        self.assertEqual(records[0].sheet, "Clinical")
        self.assertEqual(records[0].classification, "CLINICAL_SAMPLE")
        self.assertEqual(records[0].curability, "YES")
        self.assertEqual(
            warnings,
            ["source.xlsx: Could not read sheet 'Notes': damaged worksheet"],
        )

    def test_supplement_analysis_turns_reader_errors_into_not_loadable_records(
        self,
    ) -> None:
        with patch.object(
            report_curation,
            "read_supplementary_file",
            side_effect=SupplementaryParseError("cannot parse"),
        ):
            records = report_curation.analyse_supplementary_files(["/tmp/broken.xlsx"])

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].classification, "NOT_LOADABLE")
        self.assertEqual(records[0].confidence, 0)
        self.assertIn("cannot parse", records[0].verdict)
        self.assertEqual(records[0].load_error, "cannot parse")


if __name__ == "__main__":
    unittest.main()
